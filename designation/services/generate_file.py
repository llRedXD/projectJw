from ast import List
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

from designation.models import Parte, Pessoa, Reuniao


month_mapping = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}


def multi_append(lst, value, count):
    for _ in range(count):
        lst.append(value)


def get_tab1(reuniao: Reuniao, tab):
    date_obj = datetime.strptime(str(reuniao.data), "%Y-%m-%d")
    day = date_obj.day
    month = month_mapping[date_obj.month]
    formatted_date = f"{day:02d} de {month}"
    tab.append(f"{formatted_date.upper()} | {reuniao.texto.upper()}")
    tab.append("Presidente:")
    if reuniao.presidente:
        tab.append(reuniao.presidente.nome)
    else:
        tab.append("")
    tab.append("")
    tab.append("Conselheiro da sala B:")
    if reuniao.conselheiro_sala_b:
        tab.append(reuniao.conselheiro_sala_b.nome)
    else:
        tab.append("")
    tab.append("")


def get_tab2(reuniao: Reuniao, tab):
    tab.append(f"Cantico: {reuniao.cantico_inicial}")
    tab.append("Oração")
    tab.append(reuniao.oracao_inicial)
    tab.append("Comentarios Iniciais 1 (min)")


def init_tab3(tab):
    multi_append(tab, "TESOUROS DA PALAVRA DE DEUS", 2)
    tab.append("Sala B")
    tab.append("Sala Principal")


def get_tab3(parte: Parte, tab):
    nome_pessoa = parte.pessoa.nome if parte.pessoa else ""
    nome_pessoa_b = parte.pessoa_b.nome if parte.pessoa_b else ""  # Pessoa for Sala B

    if parte.numero_parte in [1, 2]:
        multi_append(
            tab,
            f"{parte.numero_parte}. {parte.nome_parte} ({parte.duracao} min)",
            3,
        )
        tab.append(f"{nome_pessoa}")
    elif parte.numero_parte == 3:
        tab.append(f"{parte.numero_parte}. {parte.nome_parte} ({parte.duracao} min)")
        tab.append("Estudante:")
        tab.append(f"{nome_pessoa_b} ({parte.ponto_parte})")
        tab.append(f"{nome_pessoa} ({parte.ponto_parte})")


def init_tab4(tab):
    tab.append("FAÇA SEU MELHOR NO MINESTERIO")
    # tab4.append("FAÇA SEU MELHOR NO MINESTERIO")
    tab.append("Sala B")
    tab.append("Sala Principal")


def get_tab4(parte: Parte, tab):
    nome_pessoa = parte.pessoa.nome if parte.pessoa else ""
    nome_pessoa_b = parte.pessoa_b.nome if parte.pessoa_b else ""  # Pessoa for Sala B
    ajudante = (
        parte.ajudante.nome if parte.ajudante else ""
    )  # Ajudante for Sala Principal
    ajudante_b = (
        parte.ajudante_b.nome if parte.ajudante_b else ""
    )  # Ajudante for Sala B
    if (
        parte.numero_parte in [4, 5, 6, 7]
        and parte.trecho == "Faça seu Melhor no Ministério"
    ):
        if parte.nome_parte != "Discurso":
            tab.append(
                f"{parte.numero_parte}. {parte.nome_parte} ({parte.duracao} min)"
            )
            tab.append(f"{nome_pessoa_b} / {ajudante_b}")
            tab.append(f"{nome_pessoa} / {ajudante}")
        else:
            tab.append(
                f"{parte.numero_parte}. {parte.nome_parte} ({parte.duracao} min)"
            )
            tab.append(f"{nome_pessoa_b}")
            tab.append(f"{nome_pessoa}")


def init_tab5(reuniao: Reuniao, tab):
    tab.append("NOSSA VIDA CRISTÃ")
    multi_append(tab, "", 2)
    multi_append(tab, f"Cantico {reuniao.cantico_meio}", 3)


def get_tab5(parte: Parte, tab):
    nome_pessoa = parte.pessoa.nome if parte.pessoa else ""
    if parte.numero_parte in [7, 8, 9] and parte.trecho == "Nossa Vida Cristã":
        multi_append(
            tab,
            f"{parte.numero_parte}. {parte.nome_parte} ({parte.duracao} min)",
            2,
        )
        tab.append(f"{nome_pessoa}")


def end_tab5(reuniao, tab):
    multi_append(tab, "Comentários finais (3 min)", 2)
    tab.append("")
    tab.append(f"Cantico: {reuniao['cantico_final']}")
    tab.append("Oração:")
    tab.append(reuniao["oracao_final_id"])


def get_tables(doc, base_index):
    tables = []
    for i in range(base_index, base_index + 5):
        table = doc.tables[i]
        if table:
            tables.append(table)
    return tables


def format_table(table, tab_data, cell_selected=None):
    count = 0
    for row in table.rows:
        for cell in row.cells:
            if cell_selected:
                if count not in cell_selected:
                    count += 1
                    continue
            for paragraph in cell.paragraphs:
                run_count = 0
                for run in paragraph.runs:
                    if run_count == 0:
                        # Save the current font properties
                        current_font_name = run.font.name
                        current_font_size = run.font.size
                        # Alter the text (modify as needed)
                        txt = str(tab_data[count])
                        if (
                            txt
                            and len(txt) > 2
                            and txt[0].isdigit()
                            and not txt[1].isdigit()
                        ):
                            # Preserve the original style for the first two characters
                            first_part = txt[:2]
                            second_part = txt[2:]
                            run.text = first_part
                            # Create a new run for the remaining text with modifications
                            new_run = paragraph.add_run(second_part)
                            new_run.font.name = current_font_name
                            new_run.font.size = Pt(11)
                            new_run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            run.text = txt
                            run.font.name = current_font_name
                            run.font.size = current_font_size
                    else:
                        run.text = ""
                    run_count += 1
            count += 1
